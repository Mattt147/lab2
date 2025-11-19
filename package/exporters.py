"""
Модуль exporters.py
Содержит классы для экспорта результатов расчётов в форматы .doc и .xls
"""

from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from .models import CalculationResult


class BaseExporter:
    """Базовый класс для экспортеров"""
    
    def __init__(self, filename=None):
        """
        Инициализация экспортера
        
        Args:
            filename (str): Имя файла для сохранения
        """
        self._filename = filename
        self._export_count = 0
    
    @property
    def filename(self):
        """Получить имя файла"""
        return self._filename
    
    @filename.setter
    def filename(self, value):
        """Установить имя файла"""
        if not value:
            raise ValueError("Имя файла не может быть пустым")
        self._filename = value
    
    def _generate_filename(self, extension):
        """
        Генерация имени файла с датой и временем
        
        Args:
            extension (str): Расширение файла
        
        Returns:
            str: Сгенерированное имя файла
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"calculation_report_{timestamp}.{extension}"
    
    def export(self, results):
        """
        Абстрактный метод экспорта (должен быть переопределён)
        
        Args:
            results: Результаты для экспорта
        """
        raise NotImplementedError("Метод export должен быть реализован в подклассе")
    
    def __str__(self):
        return f"{self.__class__.__name__}(файл: {self._filename}, экспортов: {self._export_count})"
    
    def __repr__(self):
        return f"{self.__class__.__name__}('{self._filename}')"


class DocxExporter(BaseExporter):
    """Класс для экспорта результатов в формат .docx"""
    
    def __init__(self, filename=None):
        """
        Инициализация DOCX экспортера
        
        Args:
            filename (str): Имя файла (если None, генерируется автоматически)
        """
        super().__init__(filename)
    
    def export(self, results):
        """
        Экспортировать результаты в DOCX
        
        Args:
            results (CalculationResult or list): Один результат или список результатов
        
        Returns:
            str: Путь к созданному файлу
        """
        # Преобразуем в список, если передан один результат
        if isinstance(results, CalculationResult):
            results = [results]
        
        if not results:
            raise ValueError("Нет данных для экспорта")
        
        # Генерируем имя файла, если не задано
        if not self._filename:
            self._filename = self._generate_filename("docx")
        
        # Создаём документ
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('Отчёт по расчёту отделочных материалов', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Дата создания
        date_paragraph = doc.add_paragraph()
        date_paragraph.add_run(f'Дата: {datetime.now().strftime("%d.%m.%Y %H:%M")}').italic = True
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()  # Пустая строка
        
        # Добавляем результаты
        for idx, result in enumerate(results, 1):
            # Заголовок расчёта
            doc.add_heading(f'Расчёт #{idx}: {result.material.name}', level=2)
            
            # Таблица с деталями
            table = doc.add_table(rows=5, cols=2)
            table.style = 'Light Grid Accent 1'
            
            # Заполняем таблицу
            cells = table.rows[0].cells
            cells[0].text = 'Материал:'
            cells[1].text = result.material.name
            
            cells = table.rows[1].cells
            cells[0].text = 'Площадь покрытия:'
            cells[1].text = f'{result.area:.2f} м²'
            
            cells = table.rows[2].cells
            cells[0].text = 'Запас:'
            cells[1].text = f'{result.reserve_percent}%'
            
            cells = table.rows[3].cells
            cells[0].text = 'Необходимо единиц:'
            cells[1].text = str(result.units_needed)
            
            cells = table.rows[4].cells
            cells[0].text = 'Общая стоимость:'
            cells[1].text = f'{result.total_cost:.2f} ₽'
            
            # Выделяем итоговую стоимость
            for cell in table.rows[4].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(12)
            
            doc.add_paragraph()  # Пустая строка между расчётами
        
        # Итоговая сводка
        if len(results) > 1:
            doc.add_page_break()
            doc.add_heading('Итоговая сводка', level=1)
            
            total_cost = sum(r.total_cost for r in results)
            total_area = sum(r.area for r in results)
            
            summary = doc.add_paragraph()
            summary.add_run(f'Всего расчётов: {len(results)}\n').bold = True
            summary.add_run(f'Общая площадь: {total_area:.2f} м²\n')
            summary.add_run(f'Общая стоимость: {total_cost:.2f} ₽').bold = True
        
        # Сохраняем документ
        doc.save(self._filename)
        self._export_count += 1
        
        return self._filename
    
    def __str__(self):
        return f"DocxExporter(файл: {self._filename}, экспортов: {self._export_count})"


class ExcelExporter(BaseExporter):
    """Класс для экспорта результатов в формат .xlsx"""
    
    def __init__(self, filename=None):
        """
        Инициализация Excel экспортера
        
        Args:
            filename (str): Имя файла (если None, генерируется автоматически)
        """
        super().__init__(filename)
    
    def export(self, results):
        """
        Экспортировать результаты в Excel
        
        Args:
            results (CalculationResult or list): Один результат или список результатов
        
        Returns:
            str: Путь к созданному файлу
        """
        # Преобразуем в список, если передан один результат
        if isinstance(results, CalculationResult):
            results = [results]
        
        if not results:
            raise ValueError("Нет данных для экспорта")
        
        # Генерируем имя файла, если не задано
        if not self._filename:
            self._filename = self._generate_filename("xlsx")
        
        # Создаём книгу
        wb = Workbook()
        ws = wb.active
        ws.title = "Расчёт материалов"
        
        # Стили
        header_font = Font(bold=True, size=12, color="FFFFFF")
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Заголовок отчёта
        ws.merge_cells('A1:F1')
        ws['A1'] = 'ОТЧЁТ ПО РАСЧЁТУ ОТДЕЛОЧНЫХ МАТЕРИАЛОВ'
        ws['A1'].font = Font(bold=True, size=14)
        ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
        
        # Дата
        ws.merge_cells('A2:F2')
        ws['A2'] = f'Дата: {datetime.now().strftime("%d.%m.%Y %H:%M")}'
        ws['A2'].alignment = Alignment(horizontal='right')
        
        # Заголовки таблицы
        headers = ['№', 'Материал', 'Площадь (м²)', 'Запас (%)', 'Количество', 'Стоимость (₽)']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=4, column=col)
            cell.value = header
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = border
        
        # Данные
        row = 5
        for idx, result in enumerate(results, 1):
            ws.cell(row=row, column=1, value=idx).border = border
            ws.cell(row=row, column=2, value=result.material.name).border = border
            ws.cell(row=row, column=3, value=result.area).border = border
            ws.cell(row=row, column=4, value=result.reserve_percent).border = border
            ws.cell(row=row, column=5, value=result.units_needed).border = border
            ws.cell(row=row, column=6, value=result.total_cost).border = border
            
            # Форматирование чисел
            ws.cell(row=row, column=3).number_format = '0.00'
            ws.cell(row=row, column=6).number_format = '#,##0.00 ₽'
            
            row += 1
        
        # Итоги
        if len(results) > 1:
            row += 1
            ws.cell(row=row, column=1, value='ИТОГО:').font = Font(bold=True)
            ws.cell(row=row, column=3, value=sum(r.area for r in results)).font = Font(bold=True)
            ws.cell(row=row, column=3).number_format = '0.00'
            ws.cell(row=row, column=6, value=sum(r.total_cost for r in results)).font = Font(bold=True)
            ws.cell(row=row, column=6).number_format = '#,##0.00 ₽'
        
        # Автоподбор ширины колонок
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(cell.value)
                except:
                    pass
            adjusted_width = (max_length + 2) * 1.2
            ws.column_dimensions[column_letter].width = adjusted_width
        
        # Сохраняем
        wb.save(self._filename)
        self._export_count += 1
        
        return self._filename
    
    def __str__(self):
        return f"ExcelExporter(файл: {self._filename}, экспортов: {self._export_count})"